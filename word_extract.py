from __future__ import annotations

import logging
import os
from pathlib import Path

from docx import Document
from lxml import etree

from image_convert import convert_metafile_to_png
from models import ContentItem, TestOption, TestQuestion

log = logging.getLogger(__name__)

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v": "urn:schemas-microsoft-com:vml",
    "o": "urn:schemas-microsoft-com:office:office",
}


class WordTestExtractor:
    def __init__(
            self,
            file_path: Path,
            symbol: str,
            log_small_tables: bool,
            image_output_dir: Path,
    ):
        self.file_path = Path(file_path)
        self.symbol = symbol
        self.log_small_tables = log_small_tables
        self.extract_dir = Path(image_output_dir)
        self.extract_dir.mkdir(parents=True, exist_ok=True)
        self.logs: list[str] = []  # short TK logs
        self._omml_xslt = self._load_omml_xslt()
        self._omml_xslt_missing_logged = False

    def cleanup(self) -> None:
        return None

    def _load_document(self) -> tuple[Document, Path]:
        if self.file_path.suffix.lower() != ".docx":
            raise RuntimeError("Поддерживаются только .docx")
        return Document(self.file_path), self.file_path

    # ---- Extract embedded images from docx media ----
    def _extract_images(self, doc: Document) -> dict[str, Path]:
        image_map: dict[str, Path] = {}
        count = 0
        for rel_id, part in doc.part.related_parts.items():
            if "image" not in part.content_type:
                continue
            ext = Path(part.partname).suffix
            image_path = self.extract_dir / f"{rel_id}{ext}"
            image_path.write_bytes(part.blob)
            converted_path = convert_metafile_to_png(image_path, self.extract_dir)
            image_map[rel_id] = converted_path or image_path
            count += 1
        log.info("Extracted embedded images: %d", count)
        self.logs.append(f"Изображений извлечено: {count}")
        return image_map

    def _load_omml_xslt(self) -> etree.XSLT | None:
        xslt_path = Path(__file__).with_name("omml2mml.xsl")
        if os.name == "nt":
            office_versions = ("Office16", "Office15", "Office14")
            program_files_paths = [
                Path(os.environ.get("PROGRAMFILES", r"C:\Program Files")),
                Path(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)")),
                Path(os.environ.get("PROGRAMW6432", r"C:\Program Files")),
            ]
            seen_bases: set[Path] = set()
            for base in program_files_paths:
                if base in seen_bases:
                    continue
                seen_bases.add(base)
                for version in office_versions:
                    candidates = [
                        base / "Microsoft Office" / "root" / version / "OMML2MML.XSL",
                        base / "Microsoft Office" / version / "OMML2MML.XSL",
                    ]
                    for candidate in candidates:
                        if candidate.exists():
                            try:
                                xslt_doc = etree.parse(str(candidate))
                            except Exception:
                                log.warning("Failed to parse OMML2MML XSLT at %s", candidate)
                                continue
                            log.info("Loaded OMML2MML XSLT from system path: %s", candidate)
                            return etree.XSLT(xslt_doc)

        if xslt_path.exists():
            xslt_doc = etree.parse(str(xslt_path))
            log.info("Loaded OMML2MML XSLT from local path: %s", xslt_path)
            return etree.XSLT(xslt_doc)

        log.warning("OMML2MML XSLT not found at system locations or %s", xslt_path)
        return None

    def _omml_to_mathml(self, omml_element) -> str | None:
        if omml_element is None:
            return None
        if self._omml_xslt is None:
            if not self._omml_xslt_missing_logged:
                log.warning("OMML2MML XSLT is unavailable; formulas will not be converted to MathML.")
                self._omml_xslt_missing_logged = True
            return None
        omml_xml = etree.fromstring(etree.tostring(omml_element))
        mathml = self._omml_xslt(omml_xml)
        return str(mathml)

    # ---- Parse cell content (text + images + formulas) ----
    def _content_from_cell(
            self,
            cell,
            image_map: dict[str, Path],
            formula_placeholder: str = "[formula]",
    ) -> list[ContentItem]:
        items: list[ContentItem] = []
        text_buf: list[str] = []

        def flush_text():
            if text_buf:
                items.append(ContentItem("text", "".join(text_buf)))
                text_buf.clear()

        def push_image(p: Path):
            flush_text()
            items.append(ContentItem("image", str(p)))

        def push_formula(formula_text: str | None):
            flush_text()
            items.append(
                ContentItem(
                    "formula",
                    formula_id=None,
                    path=None,
                    formula_text=formula_text,
                )
            )

        # cell children: w:p, w:tbl...
        for block in cell._tc.iterchildren():
            if not block.tag.endswith("}p"):
                continue

            # iterate direct children of paragraph in order (IMPORTANT)
            for child in list(block):
                tag = child.tag

                # OMML formula
                if tag.endswith("}oMath") or tag.endswith("}oMathPara"):
                    omml_element = child
                    if tag.endswith("}oMathPara"):
                        child_omml = child.find(".//m:oMath", namespaces=NS)
                        if child_omml is not None:
                            omml_element = child_omml
                    formula_text = self._omml_to_mathml(omml_element)
                    push_formula(formula_text)
                    continue

                # runs/hyperlinks
                if tag.endswith("}r") or tag.endswith("}hyperlink"):
                    # text
                    for t in child.findall(".//w:t", namespaces=NS):
                        if t.text:
                            text_buf.append(t.text)

                    # line breaks
                    for br in child.findall(".//w:br", namespaces=NS):
                        flush_text()
                        items.append(ContentItem("line_break"))

                    for cr in child.findall(".//w:cr", namespaces=NS):
                        flush_text()
                        items.append(ContentItem("line_break"))

                    # DrawingML images
                    for blip in child.findall(".//a:blip", namespaces=NS):
                        rid = blip.get(f"{{{NS['r']}}}embed")
                        if rid and rid in image_map:
                            push_image(image_map[rid])

                    # VML images (old equation previews)
                    for imdata in child.findall(".//v:imagedata", namespaces=NS):
                        rid = imdata.get(f"{{{NS['r']}}}id") or imdata.get(f"{{{NS['r']}}}embed")
                        if rid and rid in image_map:
                            push_image(image_map[rid])
                        else:
                            flush_text()
                            items.append(ContentItem("text", formula_placeholder))

                    # explicit OLE object marker
                    if child.find(".//o:OLEObject", namespaces=NS) is not None:
                        flush_text()
                        items.append(ContentItem("text", formula_placeholder))

            flush_text()
            items.append(ContentItem("paragraph_break"))

        while items and items[-1].item_type in {"paragraph_break", "line_break"}:
            items.pop()

        if not items:
            items.append(ContentItem("text", ""))
        return items

    def _row_has_any_content_fast(self, row) -> bool:
        """
        Fast check without consuming formula iterator.
        """
        for cell in row.cells:
            for block in cell._tc.iterchildren():
                if not block.tag.endswith("}p"):
                    continue
                if block.find(".//w:t", namespaces=NS) is not None:
                    return True
                if block.find(".//a:blip", namespaces=NS) is not None:
                    return True
                if block.find(".//v:imagedata", namespaces=NS) is not None:
                    return True
                if block.find(".//m:oMath", namespaces=NS) is not None:
                    return True
                if block.find(".//m:oMathPara", namespaces=NS) is not None:
                    return True
                if block.find(".//o:OLEObject", namespaces=NS) is not None:
                    return True
        return False

    def extract(self, formula_placeholder: str = "[formula]") -> list[TestQuestion]:
        log.info("=== EXTRACT START: %s ===", self.file_path)
        self.logs.clear()
        self.logs.append(f"Файл: {self.file_path.name}")

        doc, _ = self._load_document()
        log.info("Document loaded. Tables: %d", len(doc.tables))

        image_map = self._extract_images(doc)

        tests: list[TestQuestion] = []
        tables_used = 0

        for table_index, table in enumerate(doc.tables, start=1):
            rows = len(table.rows)
            cols = len(table.columns) if table.columns else 0
            log.debug("Table %d: rows=%d cols=%d", table_index, rows, cols)

            if rows < 3:
                if self.log_small_tables:
                    msg = f"Таблица {table_index}: < 3 строк, пропуск"
                    self.logs.append(msg)
                continue

            content_rows = sum(1 for r in table.rows if self._row_has_any_content_fast(r))
            if content_rows < 3:
                if self.log_small_tables:
                    msg = f"Таблица {table_index}: < 3 строк с контентом, пропуск"
                    self.logs.append(msg)
                continue

            tables_used += 1
            row_contents: list[list[ContentItem]] = []

            for r_i, row in enumerate(table.rows):
                row_items: list[ContentItem] = []
                for c_i, cell in enumerate(row.cells):
                    cell_items = self._content_from_cell(
                        cell,
                        image_map,
                        formula_placeholder,
                    )
                    row_items.extend(cell_items)
                row_contents.append(row_items)

            if len(row_contents) < 2:
                log.debug("Table %d: <2 content rows after parse, skip", table_index)
                continue

            question = row_contents[0]
            correct_default = row_contents[1]

            options: list[TestOption] = []
            has_marked_correct = False

            def normalize_symbol(option_items: list[ContentItem]) -> bool:
                for item in option_items:
                    if item.item_type != "text":
                        continue
                    s = item.value.lstrip()
                    if self.symbol and s.startswith(self.symbol):
                        item.value = s[len(self.symbol):].lstrip()
                        return True
                return False

            options.append(TestOption(correct_default, False))

            for option_items in row_contents[2:]:
                is_correct = False
                if self.symbol:
                    is_correct = normalize_symbol(option_items)
                    if is_correct:
                        has_marked_correct = True
                options.append(TestOption(option_items, is_correct))

            if self.symbol and normalize_symbol(correct_default):
                has_marked_correct = True
                options[0].is_correct = True

            if not has_marked_correct and options:
                options[0].is_correct = True

            for opt in options:
                if opt.is_correct:
                    correct_default = opt.content
                    break

            tests.append(TestQuestion(question=question, correct=correct_default, options=options))

            if len(tests) % 25 == 0:
                log.info("Extracted questions so far: %d", len(tests))

        log.info("Tables used: %d / %d", tables_used, len(doc.tables))
        log.info("Total tests extracted: %d", len(tests))
        self.logs.append(f"Таблиц обработано: {tables_used}")
        self.logs.append(f"Вопросов извлечено: {len(tests)}")

        log.info("=== EXTRACT END ===")
        return tests
