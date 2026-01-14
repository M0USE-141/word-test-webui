import json
import math
import os
import random
import shutil
import subprocess
import tempfile
import tkinter as tk
import uuid
from tkinter import font as tkfont
from dataclasses import dataclass, field
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

from PIL import Image, ImageTk
from docx import Document

TRANSLATIONS = {
    "ru": {
        "app_title": "Word Test Extractor",
        "main_menu": "Главное меню",
        "import_tests": "Импортировать тесты",
        "saved_tests": "Сохранённые тесты",
        "refresh_list": "Обновить список",
        "import_title": "Импорт тестов",
        "back_to_menu": "Назад в меню",
        "word_file": "Файл Word",
        "choose": "Выбрать",
        "extract_settings": "Настройки извлечения",
        "correct_symbol": "Спец. символ для правильного ответа:",
        "small_tables_log": "Показывать таблицы меньше 3 строк в логах",
        "extract_tests": "Извлечь тесты",
        "logs": "Логи",
        "test_settings_title": "Настройки теста",
        "test_not_selected": "Тест не выбран",
        "pretest_settings": "Настройки перед тестированием",
        "question_count": "Количество вопросов (0 = все):",
        "random_questions": "Случайный порядок вопросов",
        "random_options": "Случайный порядок вариантов ответов",
        "only_unanswered": "Только нерешённые вопросы",
        "show_answers_immediately": "Показывать правильный ответ сразу",
        "max_options": "Макс. вариантов ответа:",
        "start_test": "Начать тестирование",
        "testing": "Тестирование",
        "prev": "Назад",
        "next": "Дальше",
        "finish_test": "Завершить тест",
        "exit": "Выход",
        "no_questions": "Нет вопросов для тестирования.",
        "select_word": "Выберите Word файл.",
        "select_test": "Выберите тест из списка.",
        "empty_test": "Тест пустой или не найден.",
        "info": "Информация",
        "error": "Ошибка",
        "result": "Результат: {correct}/{total} правильных, отвечено {answered}, {percent:.1f}%",
        "answers": "Ответы",
        "answers_count": "Правильных ответов: {correct} из {total} ({percent:.1f}%)",
        "answer_missing": "Правильный ответ не указан.",
        "answer_correct": "Верно!",
        "answer_wrong": "Неверно. Правильный вариант: {index}",
        "extracted": "Извлечено тестов: {count}. Сохранено: {path}",
        "no_warnings": "Нет предупреждений.",
        "formula_placeholder": "[формула]",
        "questions_label": "Вопросов: {count}",
        "stats_line": "Правильно {correct}/{total} | Изучено {percent:.1f}% | Попытки: {attempts}",
        "no_attempts": "Нет попыток",
        "delete": "Удалить",
        "delete_title": "Удаление",
        "delete_confirm": "Удалить тест {name} и все связанные файлы?",
        "menu": "Меню",
        "rename": "Переименовать",
        "rename_title": "Переименование",
        "rename_prompt": "Новое название теста:",
        "rename_empty": "Название не может быть пустым.",
        "rename_exists": "Тест с таким названием уже существует.",
        "language": "Язык",
        "question_progress": "Вопрос {current} из {total}",
        "selected_test": "Выбран тест: {name}",
    },
    "uz": {
        "app_title": "Word Test Extractor",
        "main_menu": "Bosh menyu",
        "import_tests": "Testlarni import qilish",
        "saved_tests": "Saqlangan testlar",
        "refresh_list": "Roʻyxatni yangilash",
        "import_title": "Test importi",
        "back_to_menu": "Menyuga qaytish",
        "word_file": "Word fayl",
        "choose": "Tanlash",
        "extract_settings": "Ajratib olish sozlamalari",
        "correct_symbol": "Toʻgʻri javob belgisi:",
        "small_tables_log": "3 qatordan kam jadvallarni logda koʻrsatish",
        "extract_tests": "Testlarni ajratib olish",
        "logs": "Loglar",
        "test_settings_title": "Test sozlamalari",
        "test_not_selected": "Test tanlanmagan",
        "pretest_settings": "Testdan oldingi sozlamalar",
        "question_count": "Savollar soni (0 = hammasi):",
        "random_questions": "Savollar tasodifiy tartibda",
        "random_options": "Variantlar tasodifiy tartibda",
        "only_unanswered": "Faqat yechilmagan savollar",
        "show_answers_immediately": "Toʻgʻri javobni darhol koʻrsatish",
        "max_options": "Variantlar maks.:",
        "start_test": "Testni boshlash",
        "testing": "Test",
        "prev": "Orqaga",
        "next": "Keyingi",
        "finish_test": "Testni yakunlash",
        "exit": "Chiqish",
        "no_questions": "Test uchun savollar yoʻq.",
        "select_word": "Word faylni tanlang.",
        "select_test": "Roʻyxatdan testni tanlang.",
        "empty_test": "Test bo'sh yoki topilmadi.",
        "info": "Ma'lumot",
        "error": "Xatolik",
        "result": "Natija: {correct}/{total} toʻgʻri, javob berilgan {answered}, {percent:.1f}%",
        "answers": "Javoblar",
        "answers_count": "Toʻgʻri javoblar: {correct} / {total} ({percent:.1f}%)",
        "answer_missing": "Toʻgʻri javob ko'rsatilmagan.",
        "answer_correct": "Toʻgʻri!",
        "answer_wrong": "Notoʻgʻri. Toʻgʻri variant: {index}",
        "extracted": "Ajratildi: {count}. Saqlandi: {path}",
        "no_warnings": "Ogohlantirishlar yoʻq.",
        "formula_placeholder": "[formula]",
        "questions_label": "Savollar: {count}",
        "stats_line": "Toʻgʻri {correct}/{total} | Oʻrganildi {percent:.1f}% | Urinishlar: {attempts}",
        "no_attempts": "Urinishlar yoʻq",
        "delete": "Oʻchirish",
        "delete_title": "Oʻchirish",
        "delete_confirm": "{name} testini va barcha fayllarni oʻchirish?",
        "menu": "Menyu",
        "rename": "Qayta nomlash",
        "rename_title": "Qayta nomlash",
        "rename_prompt": "Testning yangi nomi:",
        "rename_empty": "Nom bo'sh bo'lishi mumkin emas.",
        "rename_exists": "Bunday nomdagi test mavjud.",
        "language": "Til",
        "question_progress": "Savol {current}/{total}",
        "selected_test": "Tanlangan test: {name}",
    },
    "en": {
        "app_title": "Word Test Extractor",
        "main_menu": "Main Menu",
        "import_tests": "Import Tests",
        "saved_tests": "Saved Tests",
        "refresh_list": "Refresh List",
        "import_title": "Test Import",
        "back_to_menu": "Back to Menu",
        "word_file": "Word File",
        "choose": "Choose",
        "extract_settings": "Extraction Settings",
        "correct_symbol": "Correct answer marker:",
        "small_tables_log": "Show tables with fewer than 3 rows in logs",
        "extract_tests": "Extract Tests",
        "logs": "Logs",
        "test_settings_title": "Test Settings",
        "test_not_selected": "No test selected",
        "pretest_settings": "Pre-test Settings",
        "question_count": "Question count (0 = all):",
        "random_questions": "Random question order",
        "random_options": "Random option order",
        "only_unanswered": "Only unanswered questions",
        "show_answers_immediately": "Show correct answer immediately",
        "max_options": "Max options:",
        "start_test": "Start Test",
        "testing": "Testing",
        "prev": "Back",
        "next": "Next",
        "finish_test": "Finish Test",
        "exit": "Exit",
        "no_questions": "No questions for testing.",
        "select_word": "Select a Word file.",
        "select_test": "Select a test from the list.",
        "empty_test": "Test is empty or not found.",
        "info": "Information",
        "error": "Error",
        "result": "Result: {correct}/{total} correct, answered {answered}, {percent:.1f}%",
        "answers": "Answers",
        "answers_count": "Correct answers: {correct} of {total} ({percent:.1f}%)",
        "answer_missing": "Correct answer is not specified.",
        "answer_correct": "Correct!",
        "answer_wrong": "Incorrect. Correct option: {index}",
        "extracted": "Extracted: {count}. Saved: {path}",
        "no_warnings": "No warnings.",
        "formula_placeholder": "[formula]",
        "questions_label": "Questions: {count}",
        "stats_line": "Correct {correct}/{total} | Learned {percent:.1f}% | Attempts: {attempts}",
        "no_attempts": "No attempts",
        "delete": "Delete",
        "delete_title": "Delete",
        "delete_confirm": "Delete test {name} and all related files?",
        "menu": "Menu",
        "rename": "Rename",
        "rename_title": "Rename",
        "rename_prompt": "New test name:",
        "rename_empty": "Name cannot be empty.",
        "rename_exists": "A test with this name already exists.",
        "language": "Language",
        "question_progress": "Question {current} of {total}",
        "selected_test": "Selected test: {name}",
    },
}


@dataclass
class ContentItem:
    item_type: str  # "text" or "image"
    value: str


@dataclass
class TestOption:
    content: list[ContentItem]
    is_correct: bool = False


@dataclass
class TestQuestion:
    question: list[ContentItem]
    correct: list[ContentItem]
    options: list[TestOption]


@dataclass
class TestSession:
    questions: list[TestQuestion]
    answers: dict[int, int] = field(default_factory=dict)
    current_index: int = 0
    option_orders: dict[int, list[TestOption]] = field(default_factory=dict)
    finished: bool = False
    answer_status: dict[int, str] = field(default_factory=dict)
    test_id: str | None = None


class WordTestExtractor:
    def __init__(
        self,
        file_path: Path,
        symbol: str,
        log_small_tables: bool,
        image_output_dir: Path,
    ):
        self.file_path = file_path
        self.symbol = symbol
        self.log_small_tables = log_small_tables
        self.extract_dir = image_output_dir
        self.logs: list[str] = []

    def cleanup(self) -> None:
        return

    def _convert_doc_to_docx(self, doc_path: Path) -> Path:
        doc_path = doc_path.resolve()
        if not doc_path.exists():
            raise RuntimeError(f"Файл не найден: {doc_path}")
        converted = doc_path.with_suffix(".docx")
        if converted.exists():
            return converted
        soffice_path = shutil.which("soffice") or shutil.which("soffice.exe")
        if not soffice_path:
            raise RuntimeError(
                "Не удалось найти soffice для конвертации .doc. "
                "Установите LibreOffice и добавьте soffice в PATH."
            )
        temp_out = Path(tempfile.mkdtemp(prefix="word_test_docx_"))
        result = subprocess.run(
            [
                soffice_path,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(temp_out),
                str(doc_path),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError(
                "Не удалось конвертировать .doc файл. "
                "Установите LibreOffice (soffice)."
            )
        candidates = sorted(temp_out.glob("*.docx"), key=lambda p: p.stat().st_mtime)
        if not candidates:
            raise RuntimeError("Конвертация .doc завершилась без результата.")
        return candidates[-1]

    def _load_document(self) -> Document:
        if self.file_path.suffix.lower() == ".doc":
            converted = self._convert_doc_to_docx(self.file_path)
            return Document(converted)
        return Document(self.file_path)

    def _extract_images(self, doc: Document) -> dict[str, Path]:
        self.extract_dir.mkdir(parents=True, exist_ok=True)
        image_map: dict[str, Path] = {}
        for rel_id, part in doc.part.related_parts.items():
            if "image" not in part.content_type:
                continue
            extension = Path(part.partname).suffix
            image_path = self.extract_dir / f"{rel_id}{extension}"
            with image_path.open("wb") as handle:
                handle.write(part.blob)
            image_map[rel_id] = image_path
        return image_map

    def _content_from_cell(self, cell, image_map: dict[str, Path]) -> list[ContentItem]:
        items: list[ContentItem] = []
        text_buffer: list[str] = []

        def flush_text() -> None:
            if text_buffer:
                items.append(ContentItem("text", "".join(text_buffer)))
                text_buffer.clear()

        for block in cell._tc.iterchildren():
            if block.tag.endswith("}p"):
                for run in block.iter():
                    if run.tag.endswith("}oMath") or run.tag.endswith("}oMathPara"):
                        flush_text()
                        math_text = "".join(run.itertext()).strip()
                        items.append(ContentItem("text", math_text or "[formula]"))
                        continue
                    if run.tag.endswith("}t") and run.text:
                        text_buffer.append(run.text)
                    if run.tag.endswith("}blip"):
                        flush_text()
                        embed = run.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if embed and embed in image_map:
                            items.append(ContentItem("image", str(image_map[embed])))
                flush_text()
        if not items:
            items.append(ContentItem("text", ""))
        return items

    def _row_has_text(self, row, image_map: dict[str, Path]) -> bool:
        for cell in row.cells:
            for item in self._content_from_cell(cell, image_map):
                if item.item_type == "text" and item.value.strip():
                    return True
                if item.item_type == "image":
                    return True
        return False

    def extract(self) -> list[TestQuestion]:
        doc = self._load_document()
        image_map = self._extract_images(doc)
        tests: list[TestQuestion] = []

        for table_index, table in enumerate(doc.tables, start=1):
            if len(table.rows) < 3:
                if self.log_small_tables:
                    self.logs.append(
                        f"Таблица {table_index}: меньше 3 строк, пропущена."
                    )
                continue
            text_rows = sum(
                1 for row in table.rows if self._row_has_text(row, image_map)
            )
            if text_rows < 3:
                if self.log_small_tables:
                    self.logs.append(
                        f"Таблица {table_index}: меньше 3 строк с текстом, пропущена."
                    )
                continue

            row_contents: list[list[ContentItem]] = []
            for row in table.rows:
                cell_items: list[ContentItem] = []
                for cell in row.cells:
                    cell_items.extend(self._content_from_cell(cell, image_map))
                row_contents.append(cell_items)

            if len(row_contents) < 2:
                continue

            question = row_contents[0]
            correct_default = row_contents[1]
            options: list[TestOption] = []
            has_marked_correct = False

            def normalize_symbol(option_items: list[ContentItem]) -> bool:
                for item in option_items:
                    if item.item_type == "text" and item.value.strip().startswith(
                        self.symbol
                    ):
                        item.value = item.value.strip()[len(self.symbol) :].lstrip()
                        return True
                return False

            options.append(TestOption(correct_default, False))

            for option_items in row_contents[2:]:
                is_correct = False
                if self.symbol:
                    is_correct = normalize_symbol(option_items)
                    if is_correct:
                        has_marked_correct = True
                options.append(TestOption(option_items, is_correct))

            if self.symbol and normalize_symbol(correct_default):
                has_marked_correct = True
                options[0].is_correct = True

            if not has_marked_correct:
                options[0].is_correct = True

            for option in options:
                if option.is_correct:
                    correct_default = option.content
                    break

            tests.append(TestQuestion(question, correct_default, options))

        return tests


class TestApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.geometry("1024x720")
        self.minsize(900, 600)
        self._apply_style()

        self.app_dir = self._get_app_data_dir()
        self.app_dir.mkdir(parents=True, exist_ok=True)
        self.wmf_cache_dir = self.app_dir / "wmf_cache"
        self.wmf_cache_dir.mkdir(parents=True, exist_ok=True)

        self.language = tk.StringVar(value="ru")
        self.title(self._t("app_title"))

        self.selected_file = tk.StringVar()
        self.symbol = tk.StringVar()
        self.log_small_tables = tk.BooleanVar(value=False)
        self.max_options = tk.IntVar(value=4)
        self.selected_test_file: Path | None = None
        self.current_test_id: str | None = None
        self.current_test_title: str | None = None

        self.question_count = tk.IntVar(value=0)
        self.random_questions = tk.BooleanVar(value=False)
        self.random_options = tk.BooleanVar(value=False)
        self.only_unanswered = tk.BooleanVar(value=False)
        self.show_answers_immediately = tk.BooleanVar(value=True)

        self.tests: list[TestQuestion] = []
        self.session: TestSession | None = None
        self.image_cache: list[ImageTk.PhotoImage] = []

        self._build_ui()

    def _build_ui(self) -> None:
        self.container = ttk.Frame(self)
        self.container.pack(fill=tk.BOTH, expand=True)

        self.main_frame = ttk.Frame(self.container, padding=10)
        self.import_frame = ttk.Frame(self.container, padding=10)
        self.settings_frame = ttk.Frame(self.container, padding=10)
        self.test_frame = ttk.Frame(self.container, padding=10)

        for frame in (self.main_frame, self.import_frame, self.settings_frame, self.test_frame):
            frame.grid(row=0, column=0, sticky="nsew")
        self.container.rowconfigure(0, weight=1)
        self.container.columnconfigure(0, weight=1)

        self._build_main_ui()
        self._build_import_ui()
        self._build_settings_ui()
        self._build_test_ui()
        self._show_frame(self.main_frame)
        self._refresh_saved_tests()

    def _show_frame(self, frame: ttk.Frame) -> None:
        frame.tkraise()

    def _t(self, key: str, **kwargs: object) -> str:
        translations = TRANSLATIONS.get(self.language.get(), TRANSLATIONS["ru"])
        template = translations.get(key, key)
        return template.format(**kwargs)

    def _rebuild_ui(self) -> None:
        self.title(self._t("app_title"))
        self.container.destroy()
        self._build_ui()
        if self.selected_test_file:
            self.selected_test_label.config(
                text=self._t("selected_test", name=self.selected_test_file.name)
            )

    def _build_main_ui(self) -> None:
        header = ttk.Label(
            self.main_frame,
            text=self._t("main_menu"),
            font=("Segoe UI", 16, "bold"),
        )
        header.pack(anchor=tk.W, pady=5)
        lang_frame = ttk.Frame(self.main_frame)
        lang_frame.pack(anchor=tk.E, pady=2)
        ttk.Label(lang_frame, text=f"{self._t('language')}:").pack(side=tk.LEFT, padx=4)
        lang_select = ttk.Combobox(
            lang_frame,
            textvariable=self.language,
            values=["ru", "uz", "en"],
            width=8,
            state="readonly",
        )
        lang_select.pack(side=tk.LEFT)
        lang_select.bind("<<ComboboxSelected>>", lambda _e: self._rebuild_ui())
        ttk.Button(
            self.main_frame, text=self._t("import_tests"), command=self._open_import
        ).pack(anchor=tk.E, pady=5)
        results_frame = ttk.LabelFrame(
            self.main_frame, text=self._t("saved_tests"), padding=10
        )
        results_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.cards_canvas = tk.Canvas(results_frame, highlightthickness=0, bg="#f7f8fa")
        self.cards_scroll = ttk.Scrollbar(
            results_frame, orient=tk.VERTICAL, command=self.cards_canvas.yview
        )
        self.cards_canvas.configure(yscrollcommand=self.cards_scroll.set)
        self.cards_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.cards_canvas.pack(fill=tk.BOTH, expand=True)
        self.cards_container = tk.Frame(self.cards_canvas, bg="#f7f8fa")
        self.cards_canvas.create_window((0, 0), window=self.cards_container, anchor="nw")
        self.cards_container.bind(
            "<Configure>",
            lambda event: self.cards_canvas.configure(
                scrollregion=self.cards_canvas.bbox("all")
            ),
        )
        ttk.Button(
            self.main_frame, text=self._t("refresh_list"), command=self._refresh_saved_tests
        ).pack(anchor=tk.E, pady=5)

    def _build_import_ui(self) -> None:
        header = ttk.Label(
            self.import_frame,
            text=self._t("import_title"),
            font=("Segoe UI", 16, "bold"),
        )
        header.pack(anchor=tk.W, pady=5)
        ttk.Button(
            self.import_frame, text=self._t("back_to_menu"), command=self._go_to_main_menu
        ).pack(anchor=tk.E)
        file_frame = ttk.LabelFrame(self.import_frame, text=self._t("word_file"), padding=10)
        file_frame.pack(fill=tk.X, pady=5)

        ttk.Entry(file_frame, textvariable=self.selected_file, width=80).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(file_frame, text=self._t("choose"), command=self._choose_file).pack(
            side=tk.LEFT, padx=5
        )

        settings_frame = ttk.LabelFrame(
            self.import_frame, text=self._t("extract_settings"), padding=10
        )
        settings_frame.pack(fill=tk.X, pady=5)

        ttk.Label(settings_frame, text=self._t("correct_symbol")).grid(
            row=0, column=0, sticky=tk.W, pady=2
        )
        ttk.Entry(settings_frame, textvariable=self.symbol, width=10).grid(
            row=0, column=1, sticky=tk.W, pady=2
        )
        ttk.Checkbutton(
            settings_frame,
            text=self._t("small_tables_log"),
            variable=self.log_small_tables,
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=2)

        ttk.Button(
            self.import_frame, text=self._t("extract_tests"), command=self._extract_tests
        ).pack(pady=10)

        self.extract_status = ttk.Label(self.import_frame, text="")
        self.extract_status.pack(anchor=tk.W)

        log_frame = ttk.LabelFrame(self.import_frame, text=self._t("logs"), padding=10)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text = tk.Text(log_frame, height=6, state=tk.DISABLED, wrap=tk.WORD)
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _build_settings_ui(self) -> None:
        header = ttk.Label(
            self.settings_frame,
            text=self._t("test_settings_title"),
            font=("Segoe UI", 16, "bold"),
        )
        header.pack(anchor=tk.W, pady=5)
        info_frame = ttk.Frame(self.settings_frame)
        info_frame.pack(fill=tk.X, pady=5)
        self.selected_test_label = ttk.Label(info_frame, text=self._t("test_not_selected"))
        self.selected_test_label.pack(side=tk.LEFT)
        ttk.Button(
            info_frame, text=self._t("back_to_menu"), command=self._go_to_main_menu
        ).pack(side=tk.RIGHT)

        settings_frame = ttk.LabelFrame(
            self.settings_frame, text=self._t("pretest_settings"), padding=10
        )
        settings_frame.pack(fill=tk.X, pady=5)

        ttk.Label(settings_frame, text=self._t("question_count")).grid(
            row=0, column=0, sticky=tk.W, pady=2
        )
        ttk.Entry(settings_frame, textvariable=self.question_count, width=10).grid(
            row=0, column=1, sticky=tk.W, pady=2
        )
        ttk.Checkbutton(
            settings_frame,
            text=self._t("random_questions"),
            variable=self.random_questions,
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=2)
        ttk.Checkbutton(
            settings_frame,
            text=self._t("random_options"),
            variable=self.random_options,
        ).grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=2)
        ttk.Checkbutton(
            settings_frame,
            text=self._t("only_unanswered"),
            variable=self.only_unanswered,
        ).grid(row=3, column=0, columnspan=2, sticky=tk.W, pady=2)
        ttk.Checkbutton(
            settings_frame,
            text=self._t("show_answers_immediately"),
            variable=self.show_answers_immediately,
        ).grid(row=4, column=0, columnspan=2, sticky=tk.W, pady=2)

        ttk.Label(settings_frame, text=self._t("max_options")).grid(
            row=0, column=2, sticky=tk.W, padx=10, pady=2
        )
        ttk.Entry(settings_frame, textvariable=self.max_options, width=5).grid(
            row=0, column=3, sticky=tk.W, pady=2
        )

        ttk.Button(
            self.settings_frame, text=self._t("start_test"), command=self._start_test
        ).pack(pady=10)

    def _build_test_ui(self) -> None:
        header = ttk.Label(
            self.test_frame,
            text=self._t("testing"),
            font=("Segoe UI", 16, "bold"),
        )
        header.pack(anchor=tk.W, pady=5)

        nav_container = ttk.Frame(self.test_frame)
        nav_container.pack(fill=tk.X, pady=5)
        self.question_nav_canvas = tk.Canvas(
            nav_container, height=40, highlightthickness=0, bg="#f7f8fa"
        )
        self.question_nav_canvas.pack(side=tk.TOP, fill=tk.X, expand=True)
        self.question_nav_scroll = ttk.Scrollbar(
            nav_container,
            orient=tk.HORIZONTAL,
            command=self.question_nav_canvas.xview,
            style="Thin.Horizontal.TScrollbar",
        )
        self.question_nav_scroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.question_nav_canvas.configure(xscrollcommand=self.question_nav_scroll.set)
        self.question_nav_frame = ttk.Frame(self.question_nav_canvas)
        self.question_nav_canvas.create_window(
            (0, 0), window=self.question_nav_frame, anchor="nw"
        )
        self.question_nav_frame.bind(
            "<Configure>",
            lambda event: self.question_nav_canvas.configure(
                scrollregion=self.question_nav_canvas.bbox("all")
            ),
        )
        self.nav_buttons: list[ttk.Button] = []

        self.question_canvas = tk.Canvas(
            self.test_frame, borderwidth=1, relief=tk.SOLID, bg="#ffffff"
        )
        self.question_scroll = ttk.Scrollbar(
            self.test_frame,
            orient=tk.VERTICAL,
            command=self.question_canvas.yview,
            style="Thin.Vertical.TScrollbar",
        )
        self.question_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.question_canvas.pack(fill=tk.BOTH, expand=True, pady=5)
        self.question_canvas.configure(yscrollcommand=self.question_scroll.set)

        self.question_container = ttk.Frame(self.question_canvas)
        self.question_canvas.create_window((0, 0), window=self.question_container, anchor="nw")
        self.question_container.bind(
            "<Configure>",
            lambda event: self.question_canvas.configure(
                scrollregion=self.question_canvas.bbox("all")
            ),
        )

        nav_buttons = ttk.Frame(self.test_frame)
        nav_buttons.pack(fill=tk.X, pady=5)
        ttk.Button(nav_buttons, text=self._t("prev"), command=self._prev_question).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(nav_buttons, text=self._t("next"), command=self._next_question).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(
            nav_buttons, text=self._t("finish_test"), command=self._finish_test
        ).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(nav_buttons, text=self._t("exit"), command=self._exit_test).pack(
            side=tk.RIGHT, padx=5
        )

        self.answer_feedback_label = ttk.Label(self.test_frame, text="")
        self.answer_feedback_label.pack(anchor=tk.W, pady=2)
        self.report_label = ttk.Label(self.test_frame, text="", foreground="blue")
        self.report_label.pack(anchor=tk.W, pady=5)

    def _apply_style(self) -> None:
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("TFrame", background="#f7f8fa")
        style.configure("TLabel", background="#f7f8fa", font=("Segoe UI", 10))
        style.configure("TButton", padding=(10, 6), font=("Segoe UI", 10))
        style.configure("TLabelframe", background="#f7f8fa", font=("Segoe UI", 10, "bold"))
        style.configure("TLabelframe.Label", background="#f7f8fa")
        style.configure("Current.TButton", background="#cfe8ff")
        style.configure("Pending.TButton", background="#e3e7ee")
        style.configure("Neutral.TButton", background="#ffe9a6")
        style.configure("Correct.TButton", background="#4caf50", foreground="white")
        style.configure("Incorrect.TButton", background="#f44336", foreground="white")
        style.configure("Thin.Vertical.TScrollbar", gripcount=0, width=6)
        style.configure("Thin.Horizontal.TScrollbar", gripcount=0, width=6)

    def _get_app_data_dir(self) -> Path:
        if os.name == "nt":
            base = Path(os.getenv("APPDATA", Path.home()))
            return base / "WordTestExtractor"
        return Path.home() / ".local" / "share" / "word_test_extractor"

    def _choose_file(self) -> None:
        file_path = filedialog.askopenfilename(
            filetypes=[("Word files", "*.doc *.docx")]
        )
        if file_path:
            self.selected_file.set(file_path)

    def _extract_tests(self) -> None:
        path = self.selected_file.get()
        if not path:
            messagebox.showwarning(self._t("error"), self._t("select_word"))
            return
        base_name = Path(path).stem
        test_id = uuid.uuid4().hex
        output_dir = self.app_dir / "extracted_tests"
        output_dir.mkdir(exist_ok=True)
        image_dir = output_dir / f"{test_id}_images"
        extractor = WordTestExtractor(
            Path(path),
            self.symbol.get().strip(),
            self.log_small_tables.get(),
            image_dir,
        )
        try:
            tests = extractor.extract()
        except Exception as exc:
            extractor.cleanup()
            messagebox.showerror(self._t("error"), str(exc))
            return

        self.tests = tests
        output_dir = self.app_dir / "extracted_tests"
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / f"{test_id}.json"
        with output_path.open("w", encoding="utf-8") as handle:
            json.dump(
                {
                    "id": test_id,
                    "title": base_name,
                    "tests": self._serialize_tests(tests),
                },
                handle,
                ensure_ascii=False,
                indent=2,
            )

        self.extract_status.config(
            text=self._t("extracted", count=len(tests), path=output_path)
        )
        self._refresh_saved_tests()
        self._update_logs(extractor.logs)
        extractor.cleanup()

    def _refresh_saved_tests(self) -> None:
        output_dir = self.app_dir / "extracted_tests"
        for widget in self.cards_container.winfo_children():
            widget.destroy()
        if not output_dir.exists():
            return
        stats = self._load_test_stats(output_dir)
        for test_file in sorted(output_dir.glob("*.json")):
            if test_file.name == "results.json":
                continue
            metadata = self._load_test_metadata(test_file)
            questions = metadata["count"]
            test_id = metadata["id"]
            title = metadata["title"]
            test_stats = stats.get(test_id, {})
            best = test_stats.get("best_score")
            attempts = test_stats.get("attempts", 0)
            correct_total = best or 0
            learned_percent = (correct_total / questions * 100) if questions else 0
            stats_line = self._t(
                "stats_line",
                correct=correct_total,
                total=questions,
                percent=learned_percent,
                attempts=attempts,
            )
            self._create_test_card(
                test_file, title, questions, stats_line, learned_percent
            )

    def _create_test_card(
        self,
        test_file: Path,
        title: str,
        questions: int,
        stats_line: str,
        learned_percent: float,
    ) -> None:
        background = self._progress_color(learned_percent)
        card = tk.Frame(
            self.cards_container,
            bg=background,
            highlightthickness=1,
            highlightbackground="#e0e0e0",
            padx=12,
            pady=10,
        )
        card.pack(fill=tk.X, pady=6)
        title = tk.Label(
            card,
            text=title,
            font=("Segoe UI", 12, "bold"),
            bg=background,
        )
        title.pack(anchor=tk.W)
        info = tk.Label(
            card,
            text=self._t("questions_label", count=questions),
            font=("Segoe UI", 10),
            bg=background,
        )
        info.pack(anchor=tk.W, pady=(2, 0))
        stats = tk.Label(
            card,
            text=stats_line,
            font=("Segoe UI", 9),
            fg="#666666",
            bg=background,
        )
        stats.pack(anchor=tk.W, pady=(2, 0))

        menu_button = tk.Button(
            card,
            text=self._t("menu"),
            bg="#e3f2fd",
            fg="#1565c0",
            relief=tk.FLAT,
            command=lambda: self._open_test_menu(test_file),
        )
        menu_button.pack(anchor=tk.E, pady=(6, 0))

        def on_click(_event: tk.Event) -> None:
            self._select_test(test_file)

        card.bind("<Button-1>", on_click)
        title.bind("<Button-1>", on_click)
        info.bind("<Button-1>", on_click)
        stats.bind("<Button-1>", on_click)

    def _progress_color(self, percent: float) -> str:
        percent = max(0.0, min(100.0, percent)) / 100.0
        start = (255, 255, 255)
        end = (190, 230, 200)
        red = int(start[0] + (end[0] - start[0]) * percent)
        green = int(start[1] + (end[1] - start[1]) * percent)
        blue = int(start[2] + (end[2] - start[2]) * percent)
        return f"#{red:02x}{green:02x}{blue:02x}"

    def _select_test(self, test_file: Path) -> None:
        self.selected_test_file = test_file
        metadata = self._load_test_metadata(test_file)
        self.selected_test_label.config(
            text=self._t("selected_test", name=metadata["title"])
        )
        self._show_frame(self.settings_frame)

    def _open_test_menu(self, test_file: Path) -> None:
        menu = tk.Toplevel(self)
        menu.title(self._t("menu"))
        menu.resizable(False, False)
        menu.grab_set()

        ttk.Label(menu, text=self._t("rename_prompt")).pack(anchor=tk.W, padx=10, pady=5)
        metadata = self._load_test_metadata(test_file)
        name_var = tk.StringVar(value=metadata["title"])
        ttk.Entry(menu, textvariable=name_var, width=40).pack(padx=10, pady=5)

        actions = ttk.Frame(menu)
        actions.pack(fill=tk.X, padx=10, pady=10)
        ttk.Button(
            actions,
            text=self._t("rename"),
            command=lambda: self._rename_test(test_file, name_var.get(), menu),
        ).pack(side=tk.LEFT)
        ttk.Button(
            actions,
            text=self._t("delete"),
            command=lambda: self._delete_test_from_menu(test_file, menu),
        ).pack(side=tk.RIGHT)

    def _rename_test(self, test_file: Path, new_name: str, menu: tk.Toplevel) -> None:
        clean_name = new_name.strip()
        if not clean_name:
            messagebox.showerror(self._t("error"), self._t("rename_empty"))
            return
        try:
            with test_file.open("r", encoding="utf-8") as handle:
                data = json.load(handle)
            if isinstance(data, list):
                data = {"id": test_file.stem, "title": clean_name, "tests": data}
            else:
                data["title"] = clean_name
            with test_file.open("w", encoding="utf-8") as handle:
                json.dump(data, handle, ensure_ascii=False, indent=2)
            if self.selected_test_file == test_file:
                self.selected_test_label.config(
                    text=self._t("selected_test", name=clean_name)
                )
        except OSError as exc:
            messagebox.showerror(self._t("error"), str(exc))
            return
        menu.destroy()
        self._refresh_saved_tests()

    def _delete_test_from_menu(self, test_file: Path, menu: tk.Toplevel) -> None:
        menu.destroy()
        self._delete_test(test_file)

    def _delete_test(self, test_file: Path) -> None:
        metadata = self._load_test_metadata(test_file)
        if not messagebox.askyesno(
            self._t("delete_title"),
            self._t("delete_confirm", name=metadata["title"]),
        ):
            return
        try:
            test_file.unlink(missing_ok=True)
            images_dir = test_file.parent / f"{metadata['id']}_images"
            shutil.rmtree(images_dir, ignore_errors=True)
            stats = self._load_test_stats(test_file.parent)
            if metadata["id"] in stats:
                stats.pop(metadata["id"], None)
                with (test_file.parent / "results.json").open(
                    "w", encoding="utf-8"
                ) as handle:
                    json.dump(stats, handle, ensure_ascii=False, indent=2)
        except OSError as exc:
            messagebox.showerror(self._t("error"), str(exc))
        self._refresh_saved_tests()

    def _load_test_stats(self, output_dir: Path) -> dict[str, dict]:
        stats_file = output_dir / "results.json"
        if not stats_file.exists():
            return {}
        with stats_file.open("r", encoding="utf-8") as handle:
            return json.load(handle)

    def _save_test_stats(self, test_file: Path, correct: int, total: int) -> None:
        output_dir = test_file.parent
        stats = self._load_test_stats(output_dir)
        test_id = self.current_test_id or test_file.stem
        record = stats.get(test_id, {})
        attempts = record.get("attempts", 0) + 1
        best_score = max(record.get("best_score", 0), correct)
        stats[test_id] = {
            "last_score": correct,
            "best_score": best_score,
            "attempts": attempts,
        }
        with (output_dir / "results.json").open("w", encoding="utf-8") as handle:
            json.dump(stats, handle, ensure_ascii=False, indent=2)

    def _count_questions(self, test_file: Path) -> int:
        with test_file.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
        if isinstance(data, list):
            return len(data)
        return len(data.get("tests", []))

    def _load_test_metadata(self, test_file: Path) -> dict[str, object]:
        with test_file.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
        if isinstance(data, list):
            return {
                "id": test_file.stem,
                "title": test_file.stem,
                "count": len(data),
            }
        return {
            "id": data.get("id", test_file.stem),
            "title": data.get("title", test_file.stem),
            "count": len(data.get("tests", [])),
        }

    def _update_logs(self, logs: list[str]) -> None:
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete("1.0", tk.END)
        if logs:
            self.log_text.insert(tk.END, "\n".join(logs))
        else:
            self.log_text.insert(tk.END, self._t("no_warnings"))
        self.log_text.config(state=tk.DISABLED)

    def _serialize_tests(self, tests: list[TestQuestion]) -> list[dict]:
        def serialize_content(content: list[ContentItem]) -> list[dict]:
            return [{"type": item.item_type, "value": item.value} for item in content]

        return [
            {
                "question": serialize_content(test.question),
                "correct": serialize_content(test.correct),
                "options": [
                    {
                        "content": serialize_content(option.content),
                        "is_correct": option.is_correct,
                    }
                    for option in test.options
                ],
            }
            for test in tests
        ]

    def _start_test(self) -> None:
        if not self.selected_test_file:
            messagebox.showwarning(self._t("error"), self._t("select_test"))
            return
        self.tests = self._load_tests_from_file(self.selected_test_file)
        if not self.tests:
            messagebox.showwarning(self._t("error"), self._t("empty_test"))
            return
        questions = list(self.tests)
        if self.only_unanswered.get():
            questions = [
                q for index, q in enumerate(questions) if index not in self._load_progress()
            ]
        if self.random_questions.get():
            random.shuffle(questions)
        count = self.question_count.get()
        if count > 0:
            questions = questions[:count]
        self.session = TestSession(questions=questions)
        self.report_label.config(text="")
        self.answer_feedback_label.config(text="")
        self._render_question_nav()
        self._show_question()
        self._show_frame(self.test_frame)

    def _load_tests_from_file(self, test_file: Path) -> list[TestQuestion]:
        with test_file.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
        if isinstance(data, list):
            tests_data = data
            test_id = test_file.stem
            title = test_file.stem
        else:
            tests_data = data.get("tests", [])
            test_id = data.get("id", test_file.stem)
            title = data.get("title", test_file.stem)
        tests: list[TestQuestion] = []
        for entry in tests_data:
            question = [ContentItem(item["type"], item["value"]) for item in entry["question"]]
            correct = [ContentItem(item["type"], item["value"]) for item in entry["correct"]]
            options = [
                TestOption(
                    [ContentItem(item["type"], item["value"]) for item in option["content"]],
                    option["is_correct"],
                )
                for option in entry["options"]
            ]
            tests.append(TestQuestion(question, correct, options))
        self.current_test_id = test_id
        self.current_test_title = title
        return tests

    def _load_progress(self) -> set[int]:
        progress_file = self.app_dir / "progress.json"
        if not progress_file.exists():
            return set()
        with progress_file.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
        return set(data.get("answered_indices", []))

    def _save_progress(self) -> None:
        if not self.session:
            return
        progress_file = self.app_dir / "progress.json"
        answered_indices = list(self.session.answers.keys())
        with progress_file.open("w", encoding="utf-8") as handle:
            json.dump({"answered_indices": answered_indices}, handle, ensure_ascii=False)

    def _render_question_nav(self) -> None:
        for widget in self.question_nav_frame.winfo_children():
            widget.destroy()
        self.nav_buttons = []
        if not self.session:
            return
        for index in range(len(self.session.questions)):
            style_name = "Pending.TButton"
            status = self.session.answer_status.get(index, "unanswered")
            if self.session.finished:
                if status == "correct":
                    style_name = "Correct.TButton"
                elif status == "incorrect":
                    style_name = "Incorrect.TButton"
            else:
                if not self.show_answers_immediately.get():
                    style_name = "Neutral.TButton"
                elif status == "correct":
                    style_name = "Correct.TButton"
                elif status == "incorrect":
                    style_name = "Incorrect.TButton"
            if index == self.session.current_index:
                style_name = "Current.TButton"
            button = ttk.Button(
                self.question_nav_frame,
                text=str(index + 1),
                width=3,
                style=style_name,
                command=lambda idx=index: self._jump_to_question(idx),
            )
            button.pack(side=tk.LEFT, padx=2)
            self.nav_buttons.append(button)
        self._center_nav_on_current()

    def _jump_to_question(self, index: int) -> None:
        if not self.session:
            return
        self.session.current_index = index
        self._show_question()
        self._center_nav_on_current()

    def _clear_question(self) -> None:
        for widget in self.question_container.winfo_children():
            widget.destroy()
        self.image_cache.clear()

    def _show_question(self) -> None:
        if not self.session:
            return
        if not self.session.questions:
            messagebox.showinfo(self._t("info"), self._t("no_questions"))
            return
        self._clear_question()
        question = self.session.questions[self.session.current_index]
        if not self.show_answers_immediately.get() or self.session.finished:
            self.answer_feedback_label.config(text="")

        ttk.Label(
            self.question_container,
            text=self._t(
                "question_progress",
                current=self.session.current_index + 1,
                total=len(self.session.questions),
            ),
            font=("Arial", 12, "bold"),
        ).pack(anchor=tk.W, pady=5)

        self._render_content_block(self.question_container, question.question)
        ttk.Separator(self.question_container, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=5)

        if self.random_options.get():
            options = self.session.option_orders.get(self.session.current_index)
            if not options:
                options = list(question.options)
                random.shuffle(options)
        else:
            options = list(question.options)

        max_opts = max(1, self.max_options.get())
        options = options[:max_opts]
        self.session.option_orders[self.session.current_index] = options

        selected_var = tk.IntVar(
            value=self.session.answers.get(self.session.current_index, -1)
        )
        selected_idx = self.session.answers.get(self.session.current_index, -1)
        correct_idx = next(
            (i for i, option in enumerate(options) if option.is_correct), None
        )

        for idx, option in enumerate(options):
            status_color = "#ffffff"
            if self.session.finished:
                if correct_idx is not None and idx == correct_idx:
                    status_color = "#c8e6c9"
                elif selected_idx == idx:
                    status_color = "#ffcdd2"
            else:
                if selected_idx == idx:
                    if not self.show_answers_immediately.get():
                        status_color = "#fff9c4"
                    elif correct_idx is not None and idx == correct_idx:
                        status_color = "#c8e6c9"
                    else:
                        status_color = "#ffcdd2"
            frame = tk.Frame(
                self.question_container,
                bg=status_color,
                highlightthickness=1,
                highlightbackground="#e0e0e0",
                padx=10,
                pady=8,
            )
            frame.pack(fill=tk.X, pady=4)
            rb = tk.Radiobutton(
                frame,
                text=f"{idx + 1}.",
                variable=selected_var,
                value=idx,
                command=lambda: self._save_answer(selected_var.get(), options),
                bg=status_color,
                anchor="w",
            )
            if self.session.finished:
                rb.config(state=tk.DISABLED)
            rb.pack(anchor=tk.W)
            self._render_content_block(frame, option.content)
        self._render_question_nav()

    def _render_content_block(self, parent, content: list[ContentItem]) -> None:
        text_font = tkfont.Font(family="Segoe UI", size=10)
        line_height = max(1, text_font.metrics("linespace"))
        max_image_height = int(line_height * 1.5)
        max_used_image_height = 0
        background = "#f5f5f5"
        if isinstance(parent, tk.Widget):
            try:
                background = parent.cget("bg")
            except tk.TclError:
                try:
                    background = parent.cget("background")
                except tk.TclError:
                    style = ttk.Style()
                    background = style.lookup(parent.winfo_class(), "background") or background
        text = tk.Text(
            parent,
            wrap=tk.WORD,
            height=1,
            borderwidth=0,
            highlightthickness=0,
            bg=background,
            font=text_font,
        )
        text.pack(fill=tk.X, anchor=tk.W)
        for item in content:
            if item.item_type == "text":
                if item.value:
                    text.insert(tk.END, item.value)
            elif item.item_type == "image" and Path(item.value).exists():
                image_path = Path(item.value)
                if image_path.suffix.lower() == ".wmf":
                    converted = self._convert_wmf_to_png(image_path)
                    if converted:
                        image_path = converted
                    else:
                        text.insert(tk.END, self._t("formula_placeholder"))
                        continue
                try:
                    image = Image.open(image_path)
                except OSError:
                    text.insert(tk.END, self._t("formula_placeholder"))
                    continue
                if image.height > max_image_height:
                    ratio = max_image_height / image.height
                    width = max(1, int(image.width * ratio))
                    image = image.resize((width, max_image_height), Image.LANCZOS)
                max_used_image_height = max(max_used_image_height, image.height)
                photo = ImageTk.PhotoImage(image)
                self.image_cache.append(photo)
                text.image_create(tk.END, image=photo)
            text.insert(tk.END, " ")
        text.update_idletasks()
        lines = int(text.index("end-1c").split(".")[0])
        image_lines = math.ceil(max_used_image_height / line_height) if max_used_image_height else 1
        text.configure(height=max(1, lines, image_lines + 1))
        text.configure(state=tk.DISABLED)

    def _convert_wmf_to_png(self, image_path: Path) -> Path | None:
        target = self.wmf_cache_dir / f"{image_path.stem}.png"
        if target.exists() and target.stat().st_mtime >= image_path.stat().st_mtime:
            return target
        soffice_path = shutil.which("soffice") or shutil.which("soffice.exe")
        if not soffice_path:
            return None
        result = subprocess.run(
            [
                soffice_path,
                "--headless",
                "--convert-to",
                "png",
                "--outdir",
                str(self.wmf_cache_dir),
                str(image_path),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode != 0:
            return None
        if target.exists():
            return target
        return None

    def _save_answer(self, selected_idx: int, options: list[TestOption]) -> None:
        if not self.session:
            return
        if self.session.finished:
            return
        self.session.answers[self.session.current_index] = selected_idx
        correct_idx = next((i for i, option in enumerate(options) if option.is_correct), None)
        if correct_idx is None:
            self.session.answer_status[self.session.current_index] = "unanswered"
        elif selected_idx == correct_idx:
            self.session.answer_status[self.session.current_index] = "correct"
        else:
            self.session.answer_status[self.session.current_index] = "incorrect"
        self._save_progress()
        self._render_question_nav()
        self._show_question()
        if self.show_answers_immediately.get():
            if correct_idx is None:
                self.answer_feedback_label.config(
                    text=self._t("answer_missing"), foreground="#ff9800"
                )
            elif selected_idx == correct_idx:
                self.answer_feedback_label.config(
                    text=self._t("answer_correct"), foreground="#4caf50"
                )
            else:
                self.answer_feedback_label.config(
                    text=self._t("answer_wrong", index=correct_idx + 1),
                    foreground="#f44336",
                )

    def _next_question(self) -> None:
        if not self.session:
            return
        if self.session.current_index < len(self.session.questions) - 1:
            self.session.current_index += 1
            self._show_question()
            self._center_nav_on_current()

    def _prev_question(self) -> None:
        if not self.session:
            return
        if self.session.current_index > 0:
            self.session.current_index -= 1
            self._show_question()
            self._center_nav_on_current()

    def _finish_test(self) -> None:
        if not self.session:
            return
        if self.session.finished:
            return
        self.session.finished = True
        total = len(self.session.questions)
        correct = 0
        answered = 0
        for idx, question in enumerate(self.session.questions):
            answer = self.session.answers.get(idx)
            if answer is None or answer == -1:
                continue
            answered += 1
            options = self.session.option_orders.get(idx, question.options)
            if answer < len(options) and options[answer].is_correct:
                correct += 1
        percent = (correct / total * 100) if total else 0
        self.report_label.config(
            text=self._t(
                "result",
                correct=correct,
                total=total,
                answered=answered,
                percent=percent,
            )
        )
        if self.selected_test_file:
            self._save_test_stats(self.selected_test_file, correct, total)
            self._refresh_saved_tests()
        if not self.show_answers_immediately.get():
            messagebox.showinfo(
                self._t("answers"),
                self._t(
                    "answers_count",
                    correct=correct,
                    total=total,
                    percent=percent,
                ),
            )
        self._render_question_nav()
        self._show_question()

    def _exit_test(self) -> None:
        self._go_to_main_menu()

    def _go_to_main_menu(self) -> None:
        self._show_frame(self.main_frame)

    def _open_import(self) -> None:
        self._show_frame(self.import_frame)

    def _center_nav_on_current(self) -> None:
        if not self.session or not self.nav_buttons:
            return
        index = self.session.current_index
        if index < 0 or index >= len(self.nav_buttons):
            return
        self.update_idletasks()
        button = self.nav_buttons[index]
        button_center = button.winfo_x() + button.winfo_width() / 2
        total_width = max(1, self.question_nav_frame.winfo_reqwidth())
        view_width = max(1, self.question_nav_canvas.winfo_width())
        target = max(0, min(button_center - view_width / 2, total_width - view_width))
        self.question_nav_canvas.xview_moveto(target / total_width)


if __name__ == "__main__":
    app = TestApp()
    app.mainloop()
