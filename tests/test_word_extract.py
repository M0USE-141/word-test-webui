from pathlib import Path

from docx import Document
from PIL import Image

from word_extract import WordTestExtractor


def test_extract_docx_with_image(tmp_path: Path) -> None:
    image_path = tmp_path / "question.png"
    Image.new("RGB", (10, 10), color="red").save(image_path)

    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    question_cell = table.cell(0, 0)
    question_cell.text = "Вопрос"
    run = question_cell.paragraphs[0].add_run()
    run.add_picture(str(image_path))

    table.cell(1, 0).text = "Правильный ответ"
    table.cell(2, 0).text = "Неправильный ответ"

    doc_path = tmp_path / "sample.docx"
    doc.save(doc_path)

    extractor = WordTestExtractor(
        file_path=doc_path,
        symbol="*",
        log_small_tables=False,
        image_output_dir=tmp_path / "assets",
    )
    results = extractor.extract()

    assert len(results) == 1
    question_items = results[0].question
    assert any(item.item_type == "image" for item in question_items)
